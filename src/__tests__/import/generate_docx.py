from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_docx(filename):
    document = Document()

    # Heading 1
    h1 = document.add_heading('Test Document', level=1)

    # Normal Paragraph
    p1 = document.add_paragraph('This is a simple paragraph. ')

    # Bold, Italic, Underline
    p1.add_run('This is bold. ').bold = True
    p1.add_run('This is italic. ').italic = True
    p1.add_run('This is underline. ').underline = True

    # Alignment
    p2 = document.add_paragraph('This paragraph is centered.')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Lists
    document.add_paragraph('Unordered item 1', style='List Bullet')
    document.add_paragraph('Unordered item 2', style='List Bullet')
    document.add_paragraph('Ordered item 1', style='List Number')
    document.add_paragraph('Ordered item 2', style='List Number')

    # Table
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Row 1, Col 1'
    table.cell(0, 1).text = 'Row 1, Col 2'
    table.cell(1, 0).text = 'Row 2, Col 1'
    table.cell(1, 1).text = 'Row 2, Col 2'

    document.save(filename)

if __name__ == '__main__':
    create_test_docx('src/__tests__/import/test.docx')
